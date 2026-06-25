#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V13.2 文件内容解析器 (File Content Parser)
解析微信文件传输助手中的文件内容（PDF/Word/Excel）

功能:
1. 解析PDF文件（使用pdfplumber/PyPDF2）
2. 解析Word文件（使用python-docx）
3. 解析Excel文件（使用openpyxl/pandas）
4. 提取文件中的文字内容
5. 识别文件中的股票代码/板块名称/财务指标
6. 保存到舆情数据库

技术路线:
- PDF: pdfplumber（更好的中文支持）
- Word: python-docx
- Excel: openpyxl + pandas
- 兜底：规则解析（仅提取文件名）

版本: V13.2
创建: 2026-06-24
作者: 毕方灵犀·天眼 (亚瑟)
"""

import os
import sys
import sqlite3
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
import json
import re

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/file_parser.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('FileParser')

class FileContentParser:
    """文件内容解析器"""
    
    def __init__(self, db_path: str = 'data/sentiment_db.db'):
        self.db_path = db_path
        
        # 解析器可用性
        self.pdf_available = False
        self.docx_available = False
        self.excel_available = False
        
        # 确保目录存在
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        
        # 检查依赖库
        self._check_dependencies()
        
        # 初始化数据库
        self._init_db()
        
        logger.info("✅ 文件内容解析器初始化完成")
    
    def _check_dependencies(self):
        """检查依赖库是否安装"""
        # 检查PDF解析库
        try:
            import pdfplumber
            self.pdf_available = True
            logger.info("✅ pdfplumber已安装（PDF解析）")
        except ImportError:
            logger.warning("⚠️ pdfplumber未安装，PDF解析功能不可用")
            logger.warning("   安装命令: pip install pdfplumber")
        
        # 检查Word解析库
        try:
            import docx
            self.docx_available = True
            logger.info("✅ python-docx已安装（Word解析）")
        except ImportError:
            logger.warning("⚠️ python-docx未安装，Word解析功能不可用")
            logger.warning("   安装命令: pip install python-docx")
        
        # 检查Excel解析库
        try:
            import openpyxl
            import pandas
            self.excel_available = True
            logger.info("✅ openpyxl/pandas已安装（Excel解析）")
        except ImportError:
            logger.warning("⚠️ openpyxl/pandas未安装，Excel解析功能不可用")
            logger.warning("   安装命令: pip install openpyxl pandas")
    
    def _init_db(self):
        """初始化数据库表"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # 创建文件解析结果表
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS file_parse_results (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT UNIQUE,
                    file_type TEXT,  -- pdf/word/excel/unknown
                    parse_engine TEXT,
                    parsed_text TEXT,
                    extracted_stocks TEXT,  -- JSON array
                    extracted_keywords TEXT,  -- JSON array
                    extracted_financials TEXT,  -- JSON object
                    confidence REAL,
                    processed INTEGER DEFAULT 0,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            conn.commit()
            conn.close()
            logger.info("✅ 数据库表初始化完成")
            
        except Exception as e:
            logger.error(f"❌ 数据库初始化失败: {e}")
            raise
    
    def parse_file(self, file_path: str) -> Dict:
        """
        解析文件内容
        
        参数:
            file_path: 文件路径
            
        返回:
            Dict: {
                'parsed_text': str,  # 解析出的文字
                'extracted_stocks': List[str],  # 提取的股票代码
                'extracted_keywords': List[str],  # 提取的关键词
                'extracted_financials': Dict,  # 提取的财务指标
                'confidence': float,  # 置信度
                'parse_engine': str,  # 使用的解析引擎
            }
        """
        logger.info(f"📄 开始解析文件: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"❌ 文件不存在: {file_path}")
            return {
                'parsed_text': '',
                'extracted_stocks': [],
                'extracted_keywords': [],
                'extracted_financials': {},
                'confidence': 0.0,
                'parse_engine': 'none',
            }
        
        # 根据文件类型选择解析方法
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            result = self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            result = self._parse_word(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            result = self._parse_excel(file_path)
        else:
            logger.warning(f"⚠️ 不支持的文件类型: {file_ext}")
            result = self._parse_with_rule(file_path)
        
        # 保存到数据库
        self._save_parse_result(file_path, result)
        
        logger.info(f"✅ 文件解析完成: {len(result['parsed_text'])} 字符, 置信度={result['confidence']:.2f}")
        return result
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """解析PDF文件"""
        logger.info("📕 解析PDF文件...")
        
        if not self.pdf_available:
            logger.warning("⚠️ pdfplumber未安装，使用规则解析")
            return self._parse_with_rule(file_path)
        
        try:
            import pdfplumber
            
            # 打开PDF文件
            with pdfplumber.open(file_path) as pdf:
                text = ""
                
                # 提取每一页的文字
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # 提取股票代码和关键词
            extracted_stocks = self._extract_stock_codes(text)
            extracted_keywords = self._extract_keywords(text)
            extracted_financials = self._extract_financials(text)
            
            result = {
                'parsed_text': text,
                'extracted_stocks': extracted_stocks,
                'extracted_keywords': extracted_keywords,
                'extracted_financials': extracted_financials,
                'confidence': 0.9,  # PDF解析置信度较高
                'parse_engine': 'pdfplumber',
            }
            
            logger.info(f"✅ PDF解析完成: {len(text)} 字符, {pdf.pages} 页")
            return result
            
        except Exception as e:
            logger.error(f"❌ PDF解析失败: {e}")
            return self._parse_with_rule(file_path)
    
    def _parse_word(self, file_path: str) -> Dict:
        """解析Word文件"""
        logger.info("📘 解析Word文件...")
        
        if not self.docx_available:
            logger.warning("⚠️ python-docx未安装，使用规则解析")
            return self._parse_with_rule(file_path)
        
        try:
            import docx
            
            # 打开Word文件
            doc = docx.Document(file_path)
            
            # 提取所有段落的文字
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # 提取股票代码和关键词
            extracted_stocks = self._extract_stock_codes(text)
            extracted_keywords = self._extract_keywords(text)
            extracted_financials = self._extract_financials(text)
            
            result = {
                'parsed_text': text,
                'extracted_stocks': extracted_stocks,
                'extracted_keywords': extracted_keywords,
                'extracted_financials': extracted_financials,
                'confidence': 0.9,  # Word解析置信度较高
                'parse_engine': 'python-docx',
            }
            
            logger.info(f"✅ Word解析完成: {len(text)} 字符")
            return result
            
        except Exception as e:
            logger.error(f"❌ Word解析失败: {e}")
            return self._parse_with_rule(file_path)
    
    def _parse_excel(self, file_path: str) -> Dict:
        """解析Excel文件"""
        logger.info("📗 解析Excel文件...")
        
        if not self.excel_available:
            logger.warning("⚠️ openpyxl/pandas未安装，使用规则解析")
            return self._parse_with_rule(file_path)
        
        try:
            import pandas as pd
            
            # 读取Excel文件（所有sheet）
            dfs = pd.read_excel(file_path, sheet_name=None)
            
            # 转换为文字
            text = ""
            for sheet_name, df in dfs.items():
                text += f"=== Sheet: {sheet_name} ===\n"
                text += df.to_string() + "\n\n"
            
            # 提取股票代码和关键词
            extracted_stocks = self._extract_stock_codes(text)
            extracted_keywords = self._extract_keywords(text)
            extracted_financials = self._extract_financials(text)
            
            result = {
                'parsed_text': text,
                'extracted_stocks': extracted_stocks,
                'extracted_keywords': extracted_keywords,
                'extracted_financials': extracted_financials,
                'confidence': 0.95,  # Excel解析置信度最高
                'parse_engine': 'pandas',
            }
            
            logger.info(f"✅ Excel解析完成: {len(dfs)} 个sheet, {len(text)} 字符")
            return result
            
        except Exception as e:
            logger.error(f"❌ Excel解析失败: {e}")
            return self._parse_with_rule(file_path)
    
    def _parse_with_rule(self, file_path: str) -> Dict:
        """
        使用规则解析文件（兜底方案）
        
        仅提取文件名中的信息
        """
        logger.info("📝 使用规则解析文件（兜底方案）...")
        
        # 从文件名提取信息
        filename = os.path.basename(file_path)
        
        # 提取股票代码（6位数字）
        stock_pattern = r'\b\d{6}\b'
        extracted_stocks = re.findall(stock_pattern, filename)
        
        # 提取关键词（简单规则）
        keywords = []
        keyword_list = ['财报', '研报', '公告', '涨停', '跌停', '买入', '卖出']
        for keyword in keyword_list:
            if keyword in filename:
                keywords.append(keyword)
        
        result = {
            'parsed_text': f'【规则解析】文件名: {filename}',
            'extracted_stocks': extracted_stocks,
            'extracted_keywords': keywords,
            'extracted_financials': {},
            'confidence': 0.2,  # 规则解析置信度较低
            'parse_engine': 'rule',
        }
        
        logger.warning("⚠️ 规则解析仅提取文件名信息，准确度较低")
        return result
    
    def _extract_stock_codes(self, text: str) -> List[str]:
        """从文本中提取股票代码（6位数字）"""
        pattern = r'\b\d{6}\b'
        stocks = re.findall(pattern, text)
        
        # 去重
        stocks = list(set(stocks))
        
        # 验证是否为有效股票代码
        valid_stocks = []
        for stock in stocks:
            if re.match(r'^(600|601|603|000|002|300)\d{3}$', stock):
                valid_stocks.append(stock)
        
        return valid_stocks
    
    def _extract_keywords(self, text: str) -> List[str]:
        """从文本中提取关键词"""
        keywords = []
        keyword_list = [
            '涨停', '跌停', '上涨', '下跌', '买入', '卖出', '持有',
            '航运', '石油', 'AI', '算力', '机器人', '无人机',
            '利好', '利空', '超预期', '不及预期',
            '财报', '研报', '公告', '业绩', '预增', '预减',
        ]
        
        for keyword in keyword_list:
            if keyword in text:
                keywords.append(keyword)
        
        return keywords
    
    def _extract_financials(self, text: str) -> Dict:
        """从文本中提取财务指标"""
        financials = {}
        
        # 提取营收（亿元）
        revenue_pattern = r'营收[：:]\s*([\d.]+)\s*亿'
        revenue_match = re.search(revenue_pattern, text)
        if revenue_match:
            financials['revenue'] = float(revenue_match.group(1))
        
        # 提取净利润（亿元）
        profit_pattern = r'净利润[：:]\s*([\d.]+)\s*亿'
        profit_match = re.search(profit_pattern, text)
        if profit_match:
            financials['profit'] = float(profit_match.group(1))
        
        # 提取增长率（%）
        growth_pattern = r'增长[：:]\s*([\d.]+)\s*%'
        growth_match = re.search(growth_pattern, text)
        if growth_match:
            financials['growth'] = float(growth_match.group(1))
        
        return financials
    
    def _save_parse_result(self, file_path: str, result: Dict):
        """保存解析结果到数据库"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
                INSERT OR REPLACE INTO file_parse_results
                (file_path, file_type, parse_engine, parsed_text, extracted_stocks, extracted_keywords, extracted_financials, confidence)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                file_path,
                os.path.splitext(file_path)[1].lower(),
                result['parse_engine'],
                result['parsed_text'],
                json.dumps(result['extracted_stocks'], ensure_ascii=False),
                json.dumps(result['extracted_keywords'], ensure_ascii=False),
                json.dumps(result['extracted_financials'], ensure_ascii=False),
                result['confidence'],
            ))
            
            conn.commit()
            conn.close()
            
            logger.info(f"✅ 解析结果已保存: {file_path}")
            
        except Exception as e:
            logger.error(f"❌ 保存解析结果失败: {e}")
    
    def batch_parse(self, file_dir: str) -> List[Dict]:
        """
        批量解析文件
        
        参数:
            file_dir: 文件目录
            
        返回:
            List[Dict]: 解析结果列表
        """
        logger.info(f"📂 开始批量解析文件: {file_dir}")
        
        if not os.path.exists(file_dir):
            logger.error(f"❌ 目录不存在: {file_dir}")
            return []
        
        # 遍历目录中的文件
        results = []
        
        for filename in os.listdir(file_dir):
            file_path = os.path.join(file_dir, filename)
            if os.path.isfile(file_path):
                result = self.parse_file(file_path)
                results.append(result)
        
        logger.info(f"✅ 批量解析完成: {len(results)} 个文件")
        return results


def main():
    """主函数"""
    logger.info("🚀 启动文件内容解析器...")
    
    # 创建解析器
    parser = FileContentParser()
    
    # 测试：解析单个文件
    test_files = [
        "data/wechat_files/test.pdf",
        "data/wechat_files/test.docx",
        "data/wechat_files/test.xlsx",
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            result = parser.parse_file(test_file)
            print(f"\n📊 解析结果 ({test_file}):")
            print(f"  解析文字: {result['parsed_text'][:100]}...")
            print(f"  股票代码: {result['extracted_stocks']}")
            print(f"  关键词: {result['extracted_keywords']}")
            print(f"  财务指标: {result['extracted_financials']}")
            print(f"  置信度: {result['confidence']:.2f}")
        else:
            logger.warning(f"⚠️ 测试文件不存在: {test_file}")


if __name__ == '__main__':
    main()
