#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V13.2 微信文件传输助手 — 文件夹实时同步系统 (WeChat File Sync)
=============================================================

基于文件系统监听的微信内容同步方案。直接监控WeChat 4.x本地文件存储目录，
实时提取新增文件内容，同步到舆情知识管道。

方案优势:
  - 零UI自动化 — 不触碰微信窗口，不依赖剪贴板/快捷键
  - 全类型支持 — PDF/DOCX/TXT/HTML/图片 自动内容提取
  - 实时+兜底 — Watchdog实时监听 + 60s定时扫描双保险
  - 不丢不重 — 文件哈希去重，只处理新文件
  - 知识管道集成 — 自动接入V13.2 KnowledgePipeline

WeChat 4.x 数据路径:
  D:/xwechat_files/zhongying6618_a734/msg/file/  — 文件传输助手文件
  D:/xwechat_files/zhongying6618_a734/msg/attach/ — 图片/附件

版本: V13.2 FS-1
创建: 2026-06-24
作者: 毕方灵犀·天眼 (亚瑟)
"""

import os
import sys
import json
import sqlite3
import gc
import time
import hashlib
import logging
import shutil
import threading
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Set
from dataclasses import dataclass, field

# Setup paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

# Configure logging
os.makedirs('logs', exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/wechat_file_sync.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('WeChatFileSync')


# ============================================================
# 配置
# ============================================================

@dataclass
class SyncConfig:
    """同步配置"""
    # WeChat 4.x 数据根目录
    wechat_data_root: str = r"D:\xwechat_files\zhongying6618_a734"
    
    # 监控的子目录（相对于 data_root）
    # msg/file/ = 文件传输助手传送的文件（PDF/DOCX/TXT等，可直接提取内容）
    # msg/attach/ = 加密图片附件(.dat)，跳过
    # msg/video/ = 视频文件，跳过
    watch_dirs: List[str] = field(default_factory=lambda: [
        "msg/file",       # 文件传输助手的文件 — 核心监控目录
    ])
    
    # 本地同步镜像目录
    sync_mirror_root: str = "data/wechat_files_sync"
    
    # 数据库路径
    db_path: str = "data/sentiment_db.db"
    
    # 支持的文件扩展名
    text_extensions: Set[str] = field(default_factory=lambda: {
        '.txt', '.md', '.html', '.htm', '.csv', '.json', '.xml', '.log', '.py', '.js', '.sh', '.bat'
    })
    pdf_extensions: Set[str] = field(default_factory=lambda: {
        '.pdf', '.PDF'
    })
    docx_extensions: Set[str] = field(default_factory=lambda: {
        '.docx', '.doc'
    })
    xlsx_extensions: Set[str] = field(default_factory=lambda: {
        '.xlsx', '.xls', '.pptx', '.ppt', '.ppsx'
    })
    image_extensions: Set[str] = field(default_factory=lambda: {
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'
    })
    
    # 扫描间隔（秒）- 作为watchdog的兜底
    scan_interval: int = 60
    
    # 最大文件大小 (bytes) - 超过此大小的文件跳过内容提取
    max_file_size: int = 100 * 1024 * 1024  # 100MB
    
    # 内容提取最大字符数
    max_content_chars: int = 50000
    
    @property
    def all_extensions(self) -> Set[str]:
        return self.text_extensions | self.pdf_extensions | self.docx_extensions | self.xlsx_extensions | self.image_extensions | self.video_extensions
    
    # 视频文件扩展名
    video_extensions: Set[str] = field(default_factory=lambda: {
        '.mp4', '.mov', '.avi', '.mkv', '.wmv', '.flv', '.m4v', '.webm'
    })
    
    # 分类阈值
    sentiment_keywords: Set[str] = field(default_factory=lambda: {
        '新闻', '事件', '政策', '市场', '行情', '涨停', '跌停', '利好', '利空',
        '突发', '公告', '预警', '风险', '危机', '制裁', '关税', '战争',
        '霍尔木兹', '中东', '美联储', '央行', '降息', '加息', 'CPI', 'GDP',
        'PMI', '就业', '失业', '通胀', '通缩', '救市', '崩盘'
    })
    knowledge_keywords: Set[str] = field(default_factory=lambda: {
        '研究', '报告', '分析', '技术', '行业', '白皮书', '深度', '策略',
        '框架', '模型', '方法论', '案例', '教程', '指南', '手册',
        '产业链', '供应链', '估值', '财报', '基本面', '财务', '调研',
        'AI', '芯片', '半导体', '新能源', '机器人', '算力', '人形机器人'
    })
    trading_keywords: Set[str] = field(default_factory=lambda: {
        '买入', '卖出', '持仓', '止损', '止盈', '仓位', '加仓', '减仓',
        '交易计划', '选股', '标的', '目标价', 'PE', 'PB', 'ROE',
        '技术面', 'K线', 'MACD', 'RSI', '均线', '量价', '突破',
        '支撑', '压力', '回踩', '反弹', '抄底'
    })


# ============================================================
# 内容智能分类引擎
# ============================================================

class ContentClassifier:
    """智能内容分类器 — 将微信文件内容自动分为舆情/知识/交易计划"""
    
    CATEGORIES = {
        'sentiment': {'label': '舆情', 'icon': '📰', 'priority': 1},
        'trading':   {'label': '交易计划', 'icon': '📊', 'priority': 2},
        'knowledge': {'label': '知识', 'icon': '📚', 'priority': 3},
        'other':     {'label': '其他', 'icon': '📎', 'priority': 9},
    }
    
    def __init__(self, config: 'SyncConfig' = None):
        self.config = config or SyncConfig()
    
    def classify(self, file_name: str, content: str, file_type: str) -> Dict:
        """
        分类文件内容
        
        Returns:
            {'category': 'sentiment'|'trading'|'knowledge'|'other',
             'label': str, 'icon': str, 'confidence': float, 'keywords_matched': [...]}
        """
        name_lower = file_name.lower()
        
        # 排除规则：文件名包含非投资类关键词 → 直接归为other
        non_trading_patterns = [
            '工资', '对账', '账单', '报销', '发票', '票据', '明细',
            '凭证', '收款', '付款', 'salary', 'invoice', 'receipt',
            '毕业', '自学考试', '自审', '成绩', '证书', '学历',
            '简历', '招聘', '求职', 'resume', 'cv',
            '海报', '宣传', '广告',
        ]
        for pat in non_trading_patterns:
            if pat in name_lower:
                return {
                    'category': 'other', 'label': '其他', 'icon': '📎',
                    'confidence': 0.9, 'keywords_matched': [f'排除:{pat}'],
                    'score_detail': {},
                }
        
        if not content or len(content.strip()) < 50:
            return self._classify_by_filename(file_name, file_type)
        
        scores = {
            'sentiment': 0,
            'trading': 0,
            'knowledge': 0,
        }
        matched = []
        
        # 关键词匹配评分
        text_lower = content.lower()
        
        for kw in self.config.sentiment_keywords:
            if kw.lower() in text_lower:
                scores['sentiment'] += 1
                matched.append(kw)
        
        for kw in self.config.trading_keywords:
            if kw.lower() in text_lower:
                scores['trading'] += 1
                matched.append(kw)
        
        for kw in self.config.knowledge_keywords:
            if kw.lower() in text_lower:
                scores['knowledge'] += 1
                matched.append(kw)
        
        # 文件名加权
        if any(k in name_lower for k in ['报告', 'report', '研究', '白皮书', '分析']):
            scores['knowledge'] += 3
        if any(k in name_lower for k in ['交易', 'trading', '持仓', '买入', '卖出']):
            scores['trading'] += 3
        if any(k in name_lower for k in ['新闻', 'news', '快讯', '突发', '事件']):
            scores['sentiment'] += 3
        
        # 文件类型加权
        ext = file_type.lower()
        if ext in ['.pdf', '.docx']:
            scores['knowledge'] += 1  # PDF/DOCX更可能是研报
        if ext in ['.html', '.htm']:
            scores['trading'] += 1   # HTML更可能是交易报告
        if ext in ['.xlsx', '.xls']:
            scores['trading'] -= 2   # Excel更多是财务数据而非交易计划
        
        # 确定分类
        best_cat = 'other'
        best_score = 0
        for cat in ['sentiment', 'trading', 'knowledge']:
            if scores[cat] > best_score:
                best_score = scores[cat]
                best_cat = cat
        
        # 计算置信度
        total_score = sum(scores.values())
        confidence = min(best_score / max(total_score, 1), 1.0) if total_score > 0 else 0.3
        
        return {
            'category': best_cat if best_score >= 1 else 'other',
            'label': self.CATEGORIES[best_cat if best_score >= 1 else 'other']['label'],
            'icon': self.CATEGORIES[best_cat if best_score >= 1 else 'other']['icon'],
            'confidence': round(confidence, 2),
            'keywords_matched': matched[:10],
            'score_detail': scores,
        }
    
    def _classify_by_filename(self, file_name: str, file_type: str) -> Dict:
        """仅根据文件名分类（内容不足时）"""
        name = file_name.lower()
        if any(k in name for k in ['交易', '持仓', '买入']):
            cat = 'trading'
        elif any(k in name for k in ['新闻', '快讯', '事件', '突发']):
            cat = 'sentiment'
        elif any(k in name for k in ['报告', '研究', '分析', '白皮书', '行业']):
            cat = 'knowledge'
        else:
            cat = 'other'
        
        return {
            'category': cat,
            'label': self.CATEGORIES[cat]['label'],
            'icon': self.CATEGORIES[cat]['icon'],
            'confidence': 0.3,
            'keywords_matched': [],
            'score_detail': {},
        }


# ============================================================
# 文件同步引擎
# ============================================================

class WeChatFileSync:
    """微信文件传输助手 — 文件夹实时同步引擎"""
    
    def __init__(self, config: SyncConfig = None):
        self.config = config or SyncConfig()
        self._running = False
        self._observer = None
        self._scan_thread = None
        
        # 验证路径
        self._validate_paths()
        
        # 初始化组件
        self.classifier = ContentClassifier(self.config)
        self._init_db()
        self._init_sync_mirror()
        self._kp_instance = None      # ★ 缓存KnowledgePipeline实例 ★
        self._ocr_reader = None      # 缓存OCR模型
        self._hash_loaded = False
        
        logger.info(f"[WeChatFileSync] 初始化完成")
        logger.info(f"[WeChatFileSync] 数据根目录: {self.config.wechat_data_root}")
        logger.info(f"[WeChatFileSync] 监控目录: {self.config.watch_dirs}")
        logger.info(f"[WeChatFileSync] 镜像目录: {self.config.sync_mirror_root}")
    
    def _validate_paths(self):
        """验证路径是否存在"""
        if not os.path.exists(self.config.wechat_data_root):
            logger.warning(f"[WeChatFileSync] WeChat数据目录不存在: {self.config.wechat_data_root}")
            logger.warning(f"[WeChatFileSync] 将在运行时持续检查…")
    
    # ============================================================
    # 数据库
    # ============================================================
    
    def _init_db(self):
        """初始化同步记录数据库"""
        try:
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS wechat_file_sync (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT NOT NULL,
                    file_name TEXT NOT NULL,
                    file_size INTEGER,
                    file_hash TEXT UNIQUE NOT NULL,
                    file_type TEXT,
                    content_extracted TEXT,
                    content_hash TEXT,
                    extracted_chars INTEGER DEFAULT 0,
                    category TEXT DEFAULT 'other',
                    category_confidence REAL DEFAULT 0,
                    keywords_matched TEXT,
                    synced_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    knowledge_pipeline_fed INTEGER DEFAULT 0,
                    sentiment_fed INTEGER DEFAULT 0
                )
            """)
            
            # 创建索引
            cursor.execute("""
                CREATE INDEX IF NOT EXISTS idx_wfs_hash 
                ON wechat_file_sync(file_hash)
            """)
            cursor.execute("""
                CREATE INDEX IF NOT EXISTS idx_wfs_synced 
                ON wechat_file_sync(synced_at)
            """)
            
            # 迁移：添加分类列（兼容旧表）
            for col, col_type in [
                ('category', 'TEXT DEFAULT "other"'),
                ('category_confidence', 'REAL DEFAULT 0'),
                ('keywords_matched', 'TEXT'),
                ('sentiment_fed', 'INTEGER DEFAULT 0'),
            ]:
                try:
                    cursor.execute(f"ALTER TABLE wechat_file_sync ADD COLUMN {col} {col_type}")
                except:
                    pass
            
            conn.commit()
            conn.close()
            logger.info("[WeChatFileSync] 数据库已初始化")
            
        except Exception as e:
            logger.error(f"[WeChatFileSync] 数据库初始化失败: {e}")
    
    def _load_processed_hashes(self):
        """加载已处理文件哈希集合"""
        try:
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT file_hash FROM wechat_file_sync")
            self._processed_hashes = {row[0] for row in cursor.fetchall()}
            conn.close()
        except Exception as e:
            logger.error(f"[WeChatFileSync] 加载哈希失败: {e}")
    
    def _init_sync_mirror(self):
        """初始化本地同步镜像目录"""
        os.makedirs(self.config.sync_mirror_root, exist_ok=True)
        # 创建按日期组织的子目录
        today_dir = os.path.join(self.config.sync_mirror_root, datetime.now().strftime('%Y-%m-%d'))
        os.makedirs(today_dir, exist_ok=True)
    
    # ============================================================
    # 文件哈希
    # ============================================================
    
    
    def _is_already_processed(self, file_hash: str) -> bool:
        """DB直查去重（不驻留内存）"""
        try:
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM wechat_file_sync WHERE file_hash = ? LIMIT 1", (file_hash,))
            row = cursor.fetchone()
            conn.close()
            return row is not None
        except Exception:
            return False

    def _mark_processed(self, file_hash: str):
        """仅标记，不驻留内存（已由DB记录）"""
        pass

    def _compute_file_hash(self, file_path: str) -> str:
        """快速文件哈希：小文件全量SHA256，大文件MD5采样"""
        file_size = os.path.getsize(file_path)
        try:
            if file_size < 5 * 1024 * 1024:  # <5MB: 全量SHA256
                sha256 = hashlib.sha256()
                with open(file_path, 'rb') as f:
                    for chunk in iter(lambda: f.read(65536), b''):
                        sha256.update(chunk)
                return sha256.hexdigest()
            else:
                # >=5MB: MD5(前64KB + 文件名 + 文件大小 + 修改时间)
                mtime = os.path.getmtime(file_path)
                name = os.path.basename(file_path)
                md5 = hashlib.md5()
                with open(file_path, 'rb') as f:
                    md5.update(f.read(65536))
                md5.update(f"{name}{file_size}{mtime}".encode())
                return md5.hexdigest()
        except Exception as e:
            logger.error(f"[WeChatFileSync] 哈希计算失败 {file_path}: {e}")
            return ""
    
    def _compute_content_hash(self, content: str) -> str:
        """计算内容哈希（用于去重）"""
        return hashlib.md5(content[:1000].encode('utf-8')).hexdigest()
    
    # ============================================================
    # 内容提取
    # ============================================================
    
    def _extract_text_file(self, file_path: str) -> Optional[str]:
        """提取文本文件内容"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc, errors='replace') as f:
                    content = f.read(self.config.max_content_chars)
                return content
            except Exception:
                continue
        return None
    
    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """提取PDF文本内容"""
        try:
            # 方案A: pypdf (快速)
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            texts = []
            total_chars = 0
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
                    total_chars += len(page_text)
                    if total_chars >= self.config.max_content_chars:
                        break
            content = '\n'.join(texts)
            if content.strip():
                return content[:self.config.max_content_chars]
        except Exception as e:
            logger.debug(f"[WeChatFileSync] pypdf提取失败: {e}")
        
        try:
            # 方案B: pdfplumber (更准确但慢)
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                texts = []
                total_chars = 0
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texts.append(page_text)
                        total_chars += len(page_text)
                        if total_chars >= self.config.max_content_chars:
                            break
                content = '\n'.join(texts)
                if content.strip():
                    return content[:self.config.max_content_chars]
        except Exception as e:
            logger.debug(f"[WeChatFileSync] pdfplumber提取失败: {e}")
        
        return None
    
    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """提取DOCX文本内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            texts = []
            total_chars = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
                    total_chars += len(para.text)
                    if total_chars >= self.config.max_content_chars:
                        break
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
                            total_chars += len(cell.text)
            content = '\n'.join(texts)
            return content[:self.config.max_content_chars] if content.strip() else None
        except Exception as e:
            logger.error(f"[WeChatFileSync] DOCX提取失败 {file_path}: {e}")
            return None
    
    def _extract_image_description(self, file_path: str) -> Optional[str]:
        """提取图片基本信息（后续可集成OCR）"""
        try:
            from PIL import Image
            img = Image.open(file_path)
            info = {
                'format': img.format,
                'size': f"{img.width}x{img.height}",
                'mode': img.mode,
                'file_size': os.path.getsize(file_path),
            }
            # 尝试OCR（如果tesseract可用）
            ocr_text = self._try_ocr(file_path)
            if ocr_text:
                info['ocr_text'] = ocr_text[:1000]
            
            return json.dumps(info, ensure_ascii=False)
        except Exception as e:
            logger.debug(f"[WeChatFileSync] 图片信息提取失败: {e}")
            return None
    
    def _try_ocr(self, image_path: str) -> Optional[str]:
        """OCR识别 — 使用easyocr (纯Python, 无需外部引擎, 中文优化)"""
        try:
            from PIL import Image
            import numpy as np
            
            img = Image.open(image_path)
            # 跳过太小的图片（头像/表情/缩略图，通常<50KB无意义）
            if os.path.getsize(image_path) < 50 * 1024:
                logger.debug(f"[WeChatFileSync] OCR跳过小图片: {os.path.basename(image_path)}")
                return None
            
            # 延迟导入easyocr（首次加载模型~2-3秒）
            import easyocr
            if not hasattr(self, '_ocr_reader'):
                self._ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
            
            # PIL → numpy (避免cv2路径编码问题)
            img_np = np.array(img)
            results = self._ocr_reader.readtext(img_np, detail=0)
            
            if results:
                text = '\n'.join(results)
                logger.info(f"[WeChatFileSync] OCR识别: {os.path.basename(image_path)} → {len(results)}行/{len(text)}字")
                return text.strip()
            return None
        except ImportError:
            logger.debug("[WeChatFileSync] easyocr未安装，跳过OCR")
            return None
        except Exception as e:
            logger.debug(f"[WeChatFileSync] OCR失败: {e}")
            return None
    
    def _extract_video_metadata(self, file_path: str) -> Optional[str]:
        """提取视频文件元数据"""
        try:
            file_size = os.path.getsize(file_path)
            mtime = datetime.fromtimestamp(os.path.getmtime(file_path))
            info = {
                'type': 'video',
                'name': os.path.basename(file_path),
                'format': os.path.splitext(file_path)[1],
                'size': file_size,
                'size_mb': round(file_size / 1024 / 1024, 2),
                'modified': mtime.strftime('%Y-%m-%d %H:%M:%S'),
            }
            # 尝试用ffprobe获取更多信息
            try:
                import subprocess
                result = subprocess.run(
                    ['ffprobe', '-v', 'quiet', '-print_format', 'json', 
                     '-show_format', '-show_streams', file_path],
                    capture_output=True, text=True, timeout=10
                )
                if result.returncode == 0:
                    probe = json.loads(result.stdout)
                    fmt = probe.get('format', {})
                    info['duration'] = fmt.get('duration', 'unknown')
                    info['bitrate'] = fmt.get('bit_rate', 'unknown')
                    for stream in probe.get('streams', []):
                        if stream.get('codec_type') == 'video':
                            info['video_codec'] = stream.get('codec_name', 'unknown')
                            info['resolution'] = f"{stream.get('width')}x{stream.get('height')}"
            except:
                pass
            
            return json.dumps(info, ensure_ascii=False)
        except Exception as e:
            logger.debug(f"[WeChatFileSync] 视频元信息提取失败: {e}")
            return None
    
    def extract_file_content(self, file_path: str) -> Optional[str]:
        """根据文件类型提取内容"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in self.config.text_extensions:
            return self._extract_text_file(file_path)
        elif ext in self.config.pdf_extensions:
            return self._extract_pdf_text(file_path)
        elif ext in self.config.docx_extensions:
            return self._extract_docx_text(file_path)
        elif ext in self.config.xlsx_extensions:
            return self._extract_office_metadata(file_path, ext)
        elif ext in self.config.image_extensions:
            return self._extract_image_description(file_path)
        elif ext in self.config.video_extensions:
            return self._extract_video_metadata(file_path)
        else:
            logger.debug(f"[WeChatFileSync] 不支持的文件类型: {ext}")
            return None
    
    def _extract_office_metadata(self, file_path: str, ext: str) -> Optional[str]:
        """提取Office文件基本信息（XLSX/PPTX等暂不解码全文本）"""
        try:
            file_size = os.path.getsize(file_path)
            mtime = datetime.fromtimestamp(os.path.getmtime(file_path))
            info = {
                'type': ext,
                'name': os.path.basename(file_path),
                'size': file_size,
                'size_mb': round(file_size / 1024 / 1024, 2),
                'modified': mtime.strftime('%Y-%m-%d %H:%M:%S'),
            }
            # 尝试用python-docx打开doc文件
            if ext == '.doc':
                try:
                    return self._extract_docx_text(file_path)
                except:
                    pass
            return json.dumps(info, ensure_ascii=False)
        except Exception as e:
            logger.debug(f"[WeChatFileSync] Office元信息提取失败: {e}")
            return None
    
    # ============================================================
    # 文件处理与同步
    # ============================================================
    
    def process_file(self, file_path: str) -> Dict:
        """
        处理单个文件：提取内容 → 保存记录 → 同步镜像 → 送入知识管道
        
        Returns:
            {'status': 'new'|'duplicate'|'skipped'|'error', ...}
        """
        result = {
            'status': 'skipped',
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_size': 0,
            'file_hash': '',
            'content_chars': 0,
        }
        
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                result['status'] = 'error'
                result['error'] = '文件不存在'
                return result
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            result['file_size'] = file_size
            if file_size > self.config.max_file_size:
                result['status'] = 'skipped'
                result['reason'] = f'文件过大 ({file_size} bytes)'
                logger.info(f"[WeChatFileSync] 跳过过大文件: {os.path.basename(file_path)} ({file_size/1024/1024:.1f}MB)")
                return result
            
            # 计算哈希
            file_hash = self._compute_file_hash(file_path)
            result['file_hash'] = file_hash
            
            # 去重检查
            if self._is_already_processed(file_hash):
                result['status'] = 'duplicate'
                return result
            
            # 提取内容
            file_type = os.path.splitext(file_path)[1].lower()
            content = self.extract_file_content(file_path)
            
            # 清理UTF-8 surrogate字符（部分PDF含emoji导致编码错误）
            if content:
                content = content.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
            
            if content:
                content_hash = self._compute_content_hash(content)
                result['content_chars'] = len(content)
            else:
                content = ''
                content_hash = ''
            
            # 🔍 智能分类
            classification = self.classifier.classify(
                os.path.basename(file_path), content, file_type
            )
            result['category'] = classification['category']
            result['category_label'] = classification['label']
            result['category_confidence'] = classification['confidence']
            
            # 保存到数据库（含分类信息）
            self._save_record(file_path, file_hash, file_type, file_size, 
                            content, content_hash, classification)
            self._mark_processed(file_hash)  # DB已记录，内存不驻留
            
            # 同步到本地镜像
            mirror_path = self._sync_to_mirror(file_path)
            result['mirror_path'] = mirror_path
            
        # 🎯 按分类路由到不同管道
            kp_count = 0
            sent_count = 0
            if content and content.strip():
                content_short = content[:3000]  # 截断, 防止大内容传参
                if classification['category'] in ['knowledge', 'trading']:
                    kp_count = self._feed_knowledge_pipeline(
                        file_path, content_short, file_type)
                if classification['category'] in ['sentiment', 'trading']:
                    sent_count = self._feed_sentiment_db(
                        file_path, content_short, file_type, classification)
                # 立即释放大字符串
                del content_short
            # 释放content (可能很大)
            if 'content' in dir():
                try:
                    del content
                except:
                    pass
            
            result['kp_count'] = kp_count
            result['sentiment_count'] = sent_count
            
            result['status'] = 'new'

            # 日志
            icon = classification['icon']
            cat_label = classification['label']
            if content and content.strip():
                preview = content[:60].replace('\n', ' ')
                logger.info(f"[WeChatFileSync] {icon} [{cat_label}] {result['file_name']} | "
                           f"{len(content)}字符 | 知识{kp_count} 舆情{sent_count} | 置信度{classification['confidence']:.0%}")
            else:
                logger.info(f"[WeChatFileSync] {icon} [{cat_label}] 新文件(无文本): {result['file_name']} | {file_type}")

        except Exception as e:
            result['status'] = 'error'
            result['error'] = str(e)
            logger.error(f"[WeChatFileSync] 处理失败 {file_path}: {e}")

        gc.collect()
        return result
    
    def _clean_surrogates(self, text: str) -> str:
        """清理UTF-8 surrogate字符，防止编码错误"""
        if not text:
            return text
        return text.encode('utf-8', errors='replace').decode('utf-8', errors='replace')

    def _save_record(self, file_path: str, file_hash: str, file_type: str,
                     file_size: int, content: str, content_hash: str,
                     classification: Dict = None):
        """保存处理记录到数据库（含分类信息）"""
        try:
            # 清理surrogate字符防止编码错误
            content = self._clean_surrogates(content)
            
            if classification is None:
                classification = {'category': 'other', 'confidence': 0, 'keywords_matched': []}
            
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO wechat_file_sync 
                (file_path, file_name, file_size, file_hash, file_type, 
                 content_extracted, content_hash, extracted_chars,
                 category, category_confidence, keywords_matched)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                file_path,
                os.path.basename(file_path),
                file_size,
                file_hash,
                file_type,
                content[:self.config.max_content_chars] if content else '',
                content_hash,
                len(content) if content else 0,
                classification.get('category', 'other'),
                classification.get('confidence', 0),
                json.dumps(classification.get('keywords_matched', []), ensure_ascii=False),
            ))
            conn.commit()
            conn.close()
        except sqlite3.IntegrityError:
            pass  # 重复哈希，正常
        except Exception as e:
            logger.error(f"[WeChatFileSync] 保存记录失败: {e}")
    
    def _sync_to_mirror(self, source_path: str) -> str:
        """同步文件到本地镜像目录"""
        try:
            file_name = os.path.basename(source_path)
            # 按日期组织
            date_dir = datetime.now().strftime('%Y-%m-%d')
            mirror_dir = os.path.join(self.config.sync_mirror_root, date_dir)
            os.makedirs(mirror_dir, exist_ok=True)
            
            # 添加时间戳避免重名
            timestamp = datetime.now().strftime('%H%M%S')
            mirror_name = f"{timestamp}_{file_name}"
            mirror_path = os.path.join(mirror_dir, mirror_name)
            
            shutil.copy2(source_path, mirror_path)
            return mirror_path
        except Exception as e:
            logger.error(f"[WeChatFileSync] 镜像同步失败: {e}")
            return ""
    
    def _feed_sentiment_db(self, file_path: str, content: str, file_type: str, 
                           classification: Dict) -> int:
        """将舆情相关内容送入舆情数据库"""
        try:
            # 清理surrogate字符
            content = self._clean_surrogates(content)
            
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            
            file_name = os.path.basename(file_path)
            content_hash = self._compute_content_hash(content)
            
            # 检查是否已存在
            cursor.execute(
                "SELECT id FROM wechat_clipboard_messages WHERE content_hash = ?",
                (content_hash,)
            )
            if cursor.fetchone():
                conn.close()
                return 0
            
            # 提取关键词
            keywords = ','.join(classification.get('keywords_matched', [])[:10])
            
            cursor.execute("""
                INSERT INTO wechat_clipboard_messages
                (sender, content, msg_time, content_hash, keywords, 
                 importance_score, sentiment_score, related_tickers)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                '微信文件传输助手',
                f"[{file_name}] {content[:500]}",
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                content_hash,
                keywords,
                70 if classification['category'] == 'sentiment' else 50,
                0.5,  # 初始情感中性
                '',   # 后续可提取股票代码
            ))
            
            conn.commit()
            
            # 更新 wechat_file_sync 记录
            cursor.execute("""
                UPDATE wechat_file_sync SET sentiment_fed = 1 
                WHERE file_path = ?
            """, (file_path,))
            conn.commit()
            
            conn.close()
            return 1
            
        except Exception as e:
            logger.error(f"[WeChatFileSync] 舆情数据库接入失败: {e}")
            return 0

    def _get_kp_instance(self):
        """获取知识管道单例实例（懒加载）"""
        if not hasattr(self, '_kp_cache'):
            self._kp_cache = None
        if self._kp_cache is not None:
            return self._kp_cache
        try:
            from V13_2_KnowledgePipeline import KnowledgePipeline
            self._kp_cache = KnowledgePipeline()
        except ImportError:
            try:
                # 容错：尝试其他常见文件名
                from knowledge_pipeline import KnowledgePipeline
                self._kp_cache = KnowledgePipeline()
            except ImportError:
                logger.debug("[WeChatFileSync] KnowledgePipeline未就绪")
                return None
        except Exception as e:
            logger.debug(f"[WeChatFileSync] KnowledgePipeline初始化失败: {e}")
            return None
        return self._kp_cache

    def _feed_knowledge_pipeline(self, file_path: str, content: str, file_type: str) -> int:
        """将提取的内容送入知识管道"""
        try:
            # 清理surrogate字符
            content = self._clean_surrogates(content)

            kp = self._get_kp_instance()
            if kp is None:
                return 0
            
            file_name = os.path.basename(file_path)
            item = {
                'title': f"微信文件: {file_name}",
                'content': content[:3000],  # 知识管道限制
                'published_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'type': f'file_{file_type.replace(".", "")}',
            }
            
            count = kp.ingest_wechat_file([item])
            return count
            
        except ImportError:
            logger.debug("[WeChatFileSync] KnowledgePipeline未就绪，跳过")
            return 0
        except Exception as e:
            logger.error(f"[WeChatFileSync] 知识管道接入失败: {e}")
            return 0
    
    # ============================================================
    # 目录扫描
    # ============================================================
    
    def _collect_files(self) -> List[str]:
        """收集所有待处理文件"""
        files = []
        for watch_dir in self.config.watch_dirs:
            full_dir = os.path.join(self.config.wechat_data_root, watch_dir)
            if not os.path.exists(full_dir):
                continue
            
            for root, dirs, filenames in os.walk(full_dir):
                for fname in filenames:
                    file_path = os.path.join(root, fname)
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in self.config.all_extensions:
                        files.append(file_path)
        
        return files
    
    def scan_and_process(self) -> Dict:
        """
        扫描所有监控目录并处理新文件
        
        Returns:
            {'scanned': int, 'new': int, 'duplicates': int, 'errors': int, 'files': [...]}
        """
        stats = {'scanned': 0, 'new': 0, 'duplicates': 0, 'skipped': 0, 'errors': 0, 'files': []}
        
        files = self._collect_files()
        stats['scanned'] = len(files)
        
        if not files:
            return stats
        
        logger.info(f"[WeChatFileSync] 扫描到 {len(files)} 个文件，检查新文件...")
        
        for file_path in files:
            result = self.process_file(file_path)
            stats[result['status']] = stats.get(result['status'], 0) + 1
            
            if result['status'] == 'new':
                stats['files'].append({
                    'name': result['file_name'],
                    'size': result['file_size'],
                    'chars': result['content_chars'],
                })
        
        if stats['new'] > 0:
            logger.info(f"[WeChatFileSync] 扫描完成: 总数{stats['scanned']} "
                       f"新增{stats['new']} 重复{stats['duplicates']} "
                       f"错误{stats['errors']}")
        
        return stats
    
    # ============================================================
    # Watchdog 实时监听
    # ============================================================
    
    def _start_watchdog(self):
        """启动watchdog文件系统监听"""
        try:
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler
            
            class WeChatFileHandler(FileSystemEventHandler):
                def __init__(self, sync_instance):
                    self.sync = sync_instance
                    self._pending = set()
                    self._lock = threading.Lock()
                
                def on_created(self, event):
                    if event.is_directory:
                        return
                    file_path = event.src_path
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext not in self.sync.config.all_extensions:
                        return
                    
                    # 去重：同一文件可能触发多次事件
                    with self._lock:
                        if file_path in self._pending:
                            return
                        self._pending.add(file_path)
                    
                    # 延迟处理：等文件写入完成
                    threading.Timer(2.0, self._delayed_process, args=[file_path]).start()
                
                def on_modified(self, event):
                    if event.is_directory:
                        return
                    file_path = event.src_path
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext not in self.sync.config.all_extensions:
                        return
                    
                    with self._lock:
                        if file_path in self._pending:
                            return
                        self._pending.add(file_path)
                    
                    threading.Timer(2.0, self._delayed_process, args=[file_path]).start()
                
                def _delayed_process(self, file_path):
                    """延迟处理文件（等待写入完成）"""
                    time.sleep(1.0)  # 额外等待
                    if os.path.exists(file_path):
                        self.sync.process_file(file_path)
                    with self._lock:
                        self._pending.discard(file_path)
            
            self._observer = Observer()
            handler = WeChatFileHandler(self)
            
            for watch_dir in self.config.watch_dirs:
                full_dir = os.path.join(self.config.wechat_data_root, watch_dir)
                if os.path.exists(full_dir):
                    self._observer.schedule(handler, full_dir, recursive=True)
                    logger.info(f"[WeChatFileSync] Watchdog 监听: {full_dir}")
                else:
                    logger.warning(f"[WeChatFileSync] Watchdog 跳过(目录不存在): {full_dir}")
            
            self._observer.start()
            logger.info("[WeChatFileSync] Watchdog 文件监听已启动")
            
        except ImportError:
            logger.warning("[WeChatFileSync] watchdog未安装，仅使用定时扫描模式")
        except Exception as e:
            logger.error(f"[WeChatFileSync] Watchdog启动失败: {e}")
    
    # ============================================================
    # 主循环
    # ============================================================
    
    def start(self, mode: str = 'continuous'):
        """
        启动同步服务
        
        Args:
            mode: 'once' 单次扫描 | 'continuous' 持续监控
        """
        self._running = True
        
        logger.info(f"[WeChatFileSync] === 启动模式: {mode} ===")
        
        if mode == 'once':
            # 单次扫描
            return self.scan_and_process()
        
        # 持续监控模式
        # 1. 先做一次全量扫描
        logger.info("[WeChatFileSync] 执行初始全量扫描...")
        self.scan_and_process()
        
        # 2. 启动watchdog
        self._start_watchdog()
        
        # 3. 定时扫描兜底
        def periodic_scan():
            while self._running:
                time.sleep(self.config.scan_interval)
                if self._running:
                    logger.debug("[WeChatFileSync] 定时扫描...")
                    self.scan_and_process()
        
        self._scan_thread = threading.Thread(target=periodic_scan, daemon=True)
        self._scan_thread.start()
        
        logger.info(f"[WeChatFileSync] 持续监控已启动 (Watchdog + {self.config.scan_interval}s兜底扫描)")
        
        # 保持主线程
        try:
            while self._running:
                time.sleep(1)
        except KeyboardInterrupt:
            self.stop()
    
    def stop(self):
        """停止同步服务"""
        logger.info("[WeChatFileSync] 正在停止...")
        self._running = False
        
        if self._observer:
            self._observer.stop()
            self._observer.join(timeout=5)
        
        logger.info("[WeChatFileSync] 已停止")
    
    # ============================================================
    # 状态与统计
    # ============================================================
    
    def get_stats(self) -> Dict:
        """获取同步统计"""
        stats = {
            'data_root': self.config.wechat_data_root,
            'data_root_exists': os.path.exists(self.config.wechat_data_root),
            'mirror_root': self.config.sync_mirror_root,
            'processed_total': 0,  # 已改为DB直查，不驻留内存processed_total不再使用内存set
            'running': self._running,
        }
        
        try:
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            
            # 今日统计
            today = datetime.now().strftime('%Y-%m-%d')
            cursor.execute("""
                SELECT COUNT(*), COALESCE(SUM(extracted_chars), 0)
                FROM wechat_file_sync
                WHERE date(synced_at) = ?
            """, (today,))
            row = cursor.fetchone()
            stats['today_files'] = row[0] or 0
            stats['today_chars'] = row[1] or 0
            
            # 总统计
            cursor.execute("SELECT COUNT(*), COALESCE(SUM(extracted_chars), 0) FROM wechat_file_sync")
            row = cursor.fetchone()
            stats['total_files'] = row[0] or 0
            stats['total_chars'] = row[1] or 0
            
            # 按类型统计
            cursor.execute("""
                SELECT file_type, COUNT(*) 
                FROM wechat_file_sync 
                GROUP BY file_type 
                ORDER BY COUNT(*) DESC
            """)
            stats['by_type'] = dict(cursor.fetchall())
            
            # 最近文件
            cursor.execute("""
                SELECT file_name, file_type, extracted_chars, synced_at
                FROM wechat_file_sync
                ORDER BY synced_at DESC
                LIMIT 10
            """)
            stats['recent'] = [
                {'name': r[0], 'type': r[1], 'chars': r[2], 'time': r[3]}
                for r in cursor.fetchall()
            ]
            
            conn.close()
        except Exception as e:
            stats['db_error'] = str(e)
        
        return stats


# ============================================================
# 命令行接口
# ============================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description='V13.2 微信文件传输助手 — 文件夹实时同步系统'
    )
    parser.add_argument('--once', action='store_true', help='单次扫描后退出')
    parser.add_argument('--watch', action='store_true', default=True, help='持续监控模式(默认)')
    parser.add_argument('--stats', action='store_true', help='显示统计信息')
    parser.add_argument('--scan-interval', type=int, default=60, help='定时扫描间隔(秒)')
    parser.add_argument('--test', action='store_true', help='测试模式：扫描并显示结果')
    args = parser.parse_args()
    
    config = SyncConfig()
    if args.scan_interval != 60:
        config.scan_interval = args.scan_interval
    
    sync = WeChatFileSync(config)
    
    if args.stats:
        stats = sync.get_stats()
        print(json.dumps(stats, indent=2, ensure_ascii=False))
        return
    
    if args.once or args.test:
        logger.info("[WeChatFileSync] 执行单次扫描...")
        result = sync.scan_and_process()
        print(json.dumps(result, indent=2, ensure_ascii=False))
        return
    
    # 持续监控
    sync.start(mode='continuous')


if __name__ == '__main__':
    main()
